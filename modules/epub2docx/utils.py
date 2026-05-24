#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
EPUB转Word - DOCX XML 级后处理工具集
融合版本：结合参考实现的核心成功技术 + 我们的框架

v3.2.0 新增：5个精选排版预设（学术论文/书籍排版/商务报告/技术文档/保留原样）
"""

import logging
import copy
from pathlib import Path
from typing import Optional, Callable

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from lxml import etree

logger = logging.getLogger(__name__)


# ==================== 5个精选排版预设 ====================

DOCX_PRESETS = {
    
    # ========== 预设1：书籍排版 ==========
    "book": {
        "name": "📖 书籍排版",
        "body": {
            "cn_font": "SimSun",           # 正文宋体
            "en_font": "Times New Roman",  # 学术英文标准
            "size_pt": 12.0,               # 12pt，小四
            "line_spacing": 1.5,           # 1.5倍行距
            "indent_chars": 0              # 缩进0字符
        },
        "heading": {
            "cn_font": "KaiTi",            # 标题楷体
            "en_font": "Times New Roman",  # 学术英文标准
            "base_size_pt": 14.0,          # 一级标题：14pt（四号）
            "is_bold": True                # 标题加粗
        }
    },

    # ========== 预设2：学术论文 ==========
    "academic": {
        "name": "📚 学术论文",
        "body": {
            "cn_font": "SimSun",           # 正文宋体
            "en_font": "Times New Roman",  # 学术英文标准
            "size_pt": 10.5,               # 10.5pt，五号
            "line_spacing": 1.5,           # 1.5倍行距
            "indent_chars": 1.0            # 缩进1字符
        },
        "heading": {
            "cn_font": "SimHei",           # 标题黑体
            "en_font": "Times New Roman",  # 学术英文标准
            "base_size_pt": 14.0,          # 一级标题：14pt（四号）
            "is_bold": True                # 标题加粗
        }
    },

    # ========== 预设3：技术文档 ==========
    "technical": {
        "name": "🔧 技术文档",
        "body": {
            "cn_font": "SimSun",           # 正文宋体
            "en_font": "Consolas",         # 等宽字体（适合代码）
            "size_pt": 9.0,                # 9.0pt，小五
            "line_spacing": 1.25,          # 1.25倍行距
            "indent_chars": 0              # 缩进0字符
        },
        "heading": {
            "cn_font": "SimHei",           # 标题黑体
            "en_font": "Consolas",         # 等宽字体（适合代码）
            "base_size_pt": 12.0,          # 一级标题：12pt（小四）
            "is_bold": True                # 标题加粗
        }
    },    

    # ========== 预设4：商务报告 ==========
    "business": {
        "name": "💼 商务报告",
        "body": {
            "cn_font": "Microsoft YaHei",  # 正文：微软雅黑
            "en_font": "Arial",            # Arial（无衬线）
            "size_pt": 10.5,               # 10.5pt，五号
            "line_spacing": 1.25,          # 1.25倍行距
            "indent_chars": 0              # 缩进0字符
        },
        "heading": {
            "cn_font": "Microsoft YaHei",  # 标题：微软雅黑
            "en_font": "Arial",            # Arial（无衬线）
            "base_size_pt": 14.0,          # 一级标题：14pt（四号）
            "is_bold": True                # 标题加粗
        }
    },

    # ========== 预设5：保留原样 ==========
    "none": {
        "name": "⏸️ 保留原样"
    }
}

# ==================== XML 辅助函数 ====================

def _get_or_create_child(parent, tag_str):
    """获取或创建 XML 子节点"""
    child = parent.find(qn(tag_str))
    if child is None:
        child = OxmlElement(tag_str)
        parent.append(child)
    return child


def _force_set_fonts(rPr_element, cn_font: str, en_font: str):
    """【绝杀】强制设置字体，彻底清空 Calibre 的主题引用"""
    rFonts = rPr_element.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr_element.insert(0, rFonts)
    
    # ★ 关键：清空所有原有属性（包括 asciiTheme, eastAsiaTheme, hint 等）
    rFonts.attrib.clear()
    
    # 重新设置绝对字体名称
    rFonts.set(qn('w:ascii'), en_font)      # 英文/数字
    rFonts.set(qn('w:hAnsi'), en_font)      # 高 ANSI
    rFonts.set(qn('w:cs'), en_font)         # 复杂脚本
    rFonts.set(qn('w:eastAsia'), cn_font)   # 东亚字体（中文）
    rFonts.set(qn('w:hint'), 'eastAsia')    # 提示使用东亚字体


def _set_paragraph_format(style, line_spacing: float = 1.5, indent_chars: float = 0.0,
                          space_before_pt: float = 0, space_after_pt: float = 0):
    """设置段落格式（缩进、行距、段间距）"""
    pPr = style.element.get_or_add_pPr()
    
    spacing = _get_or_create_child(pPr, 'w:spacing')
    spacing.set(qn('w:line'), str(int(240 * line_spacing)))
    spacing.set(qn('w:lineRule'), 'auto')
    if space_before_pt > 0:
        spacing.set(qn('w:before'), str(int(space_before_pt * 20)))
    if space_after_pt > 0:
        spacing.set(qn('w:after'), str(int(space_after_pt * 20)))
    
    if indent_chars > 0:
        ind = _get_or_create_child(pPr, 'w:ind')
        ind.set(qn('w:firstLineChars'), str(int(indent_chars * 100)))
    else:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)


def _clean_all_font_sizes(doc: Document, log_callback: Optional[Callable] = None) -> int:
    """强制删除所有运行上的直接字号设置（w:sz 和 w:szCs）
    
    这是实现精确字号控制的关键步骤。
    删除后，字号将由样式定义或 docDefaults 决定。
    
    Args:
        doc: python-docx Document 对象
        log_callback: 日志回调函数
        
    Returns:
        int: 清洗的字号节点数量
    """
    cleaned = 0
    
    # 遍历所有段落
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._r.find(qn('w:rPr'))
            if rPr is not None:
                # 删除 w:sz 节点（字号）
                sz_nodes = rPr.findall(qn('w:sz'))
                for node in sz_nodes:
                    rPr.remove(node)
                    cleaned += 1
                
                # 删除 w:szCs 节点（复杂脚本字号）
                szCs_nodes = rPr.findall(qn('w:szCs'))
                for node in szCs_nodes:
                    rPr.remove(node)
                    cleaned += 1
    
    # 遍历所有表格中的单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        rPr = run._r.find(qn('w:rPr'))
                        if rPr is not None:
                            sz_nodes = rPr.findall(qn('w:sz'))
                            for node in sz_nodes:
                                rPr.remove(node)
                                cleaned += 1
                            szCs_nodes = rPr.findall(qn('w:szCs'))
                            for node in szCs_nodes:
                                rPr.remove(node)
                                cleaned += 1
    
    if log_callback and cleaned > 0:
        log_callback(f"   🧹 清洗字号: {cleaned} 处")
    
    return cleaned


def _clean_run_direct_formatting(run):
    """清洗 Calibre 遗留的 Run 级别直接格式（字体 + 字号）"""
    rPr = run._r.find(qn('w:rPr'))
    if rPr is not None:
        # 删除字体节点
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            rPr.remove(rFonts)
        # 删除字号节点
        for tag in [qn('w:sz'), qn('w:szCs')]:
            for node in rPr.findall(tag):
                rPr.remove(node)


# ==================== 核心排版引擎 ====================

def apply_style_preset(
    docx_path: Path,
    preset_key: str,
    log_callback: Optional[Callable] = None
) -> bool:
    """应用排版预设 - 含字号清洗功能
    
    v3.2.0 支持5个精选预设：academic, book, business, technical, none
    """
    if preset_key == "none" or preset_key not in DOCX_PRESETS:
        if log_callback:
            log_callback(f"⏸️ 排版预设: 保留原样")
        return True
    
    preset = DOCX_PRESETS[preset_key]
    body_cfg = preset["body"]
    head_cfg = preset["heading"]
    
    try:
        if log_callback:
            log_callback(f"🎨 应用排版预设: {preset['name']}")
            log_callback(f"   正文: {body_cfg['cn_font']} {body_cfg['size_pt']}pt")
            log_callback(f"   标题: {head_cfg['cn_font']}")
        
        doc = Document(docx_path)
        
        # ===== 阶段0：强制清洗所有直接字号设置 =====
        cleaned_sizes = _clean_all_font_sizes(doc, log_callback)
        
        # ===== 阶段1：设置 docDefaults（全局默认字体）=====
        styles_element = doc.styles.element
        doc_defaults = styles_element.find(qn('w:docDefaults'))
        if doc_defaults is None:
            doc_defaults = OxmlElement('w:docDefaults')
            styles_element.insert(0, doc_defaults)
        
        rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            rPrDefault = OxmlElement('w:rPrDefault')
            doc_defaults.append(rPrDefault)
        
        rPr = rPrDefault.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            rPrDefault.append(rPr)
        
        # 设置全局默认字体
        _force_set_fonts(rPr, body_cfg['cn_font'], body_cfg['en_font'])
        
        # 设置全局默认字号
        sz = _get_or_create_child(rPr, 'w:sz')
        sz.set(qn('w:val'), str(int(body_cfg['size_pt'] * 2)))
        szCs = _get_or_create_child(rPr, 'w:szCs')
        szCs.set(qn('w:val'), str(int(body_cfg['size_pt'] * 2)))
        
        # ===== 阶段2：遍历并修改所有样式 =====
        for style in doc.styles:
            # 确保 rPr 存在
            style_rPr = style.element.find(qn('w:rPr'))
            if style_rPr is None:
                style_rPr = OxmlElement('w:rPr')
                style.element.append(style_rPr)
            
            style_name = style.name.lower()
            is_heading = style_name.startswith('heading') or 'heading' in style_name
            
            if is_heading:
                _force_set_fonts(style_rPr, head_cfg['cn_font'], head_cfg['en_font'])
                # 设置标题字号
                sz_heading = _get_or_create_child(style_rPr, 'w:sz')
                sz_heading.set(qn('w:val'), str(int(head_cfg['base_size_pt'] * 2)))
                if head_cfg.get('is_bold'):
                    b = style_rPr.find(qn('w:b'))
                    if b is None:
                        style_rPr.append(OxmlElement('w:b'))
            else:
                _force_set_fonts(style_rPr, body_cfg['cn_font'], body_cfg['en_font'])
                # 设置正文字号
                sz_body = _get_or_create_child(style_rPr, 'w:sz')
                sz_body.set(qn('w:val'), str(int(body_cfg['size_pt'] * 2)))
        
        # ===== 阶段3：设置段落格式 =====
        if 'Normal' in doc.styles:
            _set_paragraph_format(
                doc.styles['Normal'],
                line_spacing=body_cfg.get('line_spacing', 1.5),
                indent_chars=body_cfg.get('indent_chars', 0.0)
            )
        
        # 设置标题段落格式
        for i in range(1, 7):
            style_name = f'Heading {i}'
            if style_name in doc.styles:
                space_before = 12.0 if i <= 2 else 6.0
                _set_paragraph_format(
                    doc.styles[style_name],
                    line_spacing=1.2,
                    space_before_pt=space_before,
                    space_after_pt=6.0
                )
        
        # ===== 阶段4：清洗残留的直接格式（字体） =====
        cleaned_fonts = 0
        for para in doc.paragraphs:
            for run in para.runs:
                rPr = run._r.find(qn('w:rPr'))
                if rPr is not None and rPr.find(qn('w:rFonts')) is not None:
                    _clean_run_direct_formatting(run)
                    cleaned_fonts += 1
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            rPr = run._r.find(qn('w:rPr'))
                            if rPr is not None and rPr.find(qn('w:rFonts')) is not None:
                                _clean_run_direct_formatting(run)
                                cleaned_fonts += 1
        
        doc.save(docx_path)
        
        if log_callback:
            log_callback(f"✅ 排版完成: 清洗字号 {cleaned_sizes} 处, 清洗字体 {cleaned_fonts} 处")
        
        return True
        
    except Exception as e:
        if log_callback:
            log_callback(f"❌ 排版失败: {e}")
        logger.error(f"排版预设应用失败: {e}")
        return False


# ==================== 软回车修复 ====================

def fix_soft_line_breaks(docx_path: Path, log_callback: Optional[Callable] = None) -> bool:
    """修复软回车：将 Word 中的软回车转换为硬段落"""
    try:
        if log_callback:
            log_callback("🔧 开始转换软回车为硬段落...")
        doc = Document(docx_path)
        modified_count = 0

        for paragraph in doc.paragraphs:
            p_elem = paragraph._element
            br_nodes = p_elem.findall(f'.//{qn("w:br")}')
            if not br_nodes:
                continue

            parent_tag = p_elem.getparent().tag
            is_in_body = parent_tag == qn('w:body')

            for br in reversed(br_nodes):
                run = br.getparent()
                if run is None or run.tag != qn('w:r'):
                    continue
                if is_in_body:
                    _split_paragraph_at_br(p_elem, br, run)
                else:
                    run.remove(br)
                modified_count += 1

        if modified_count > 0:
            doc.save(docx_path)
            if log_callback:
                log_callback(f"✅ 处理 {modified_count} 处换行符")
            return True
        return False
    except Exception as e:
        if log_callback:
            log_callback(f"⚠️ 软回车修复失败: {e}")
        return False


def _split_paragraph_at_br(p_elem, br, run):
    """在软回车处拆分段落"""
    new_p = etree.SubElement(p_elem.getparent(), qn('w:p'))
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is not None:
        new_pPr = copy.deepcopy(pPr)
        numPr = new_pPr.find(qn('w:numPr'))
        if numPr is not None:
            new_pPr.remove(numPr)
        new_p.append(new_pPr)

    new_r = etree.SubElement(new_p, qn('w:r'))
    rPr = run.find(qn('w:rPr'))
    if rPr is not None:
        new_r.append(copy.deepcopy(rPr))

    run_children = list(run)
    try:
        idx = run_children.index(br)
        for child in run_children[idx + 1:]:
            new_r.append(child)
    except ValueError:
        pass

    p_children = list(p_elem)
    try:
        r_idx = p_children.index(run)
        for r in p_children[r_idx + 1:]:
            if r.tag == qn('w:r'):
                new_p.append(r)
    except ValueError:
        pass

    if br.getparent() is not None:
        br.getparent().remove(br)
    p_elem.addnext(new_p)


# ==================== 向后兼容 ====================

force_apply_style_preset = apply_style_preset


# ==================== 元信息 ====================
__author__ = "YQJ"
__version__ = "3.2.0"
__date__ = "2026.05.24"